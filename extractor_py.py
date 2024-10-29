# %%
import pandas as pd
from docx import Document
import os
import glob
import win32com.client as win32

# %%
# MUDAR QUANDO EXPORTO
# Get the directory where the current Python file is located

try:
    current_directory = __file__
    current_directory = current_directory.rsplit('\\', 1)[0]
except:
    current_directory = os.path.dirname(os.path.abspath("extrator.ipynb"))

# %%
# Change the current working directory to that directory
os.chdir(current_directory)

# %%
# Define the pattern to match the filename
file_pattern = 'Consulta de Equilíbrio(*).xlsx'

# Use glob to find the file
files = glob.glob(file_pattern)

# Check if any files were found
if files:
    # Assuming you want to use the first matching file
    excel_file = files[0]
    
    # Load the Excel data
    df = pd.read_excel(excel_file)
    df['Processo'] = pd.to_numeric(df['Processo'], errors='coerce').astype('Int64')  # Using 'Int64' to handle NaN values

    # Print the name of the file that was opened
    print(f"Opened file: {excel_file}")
else:
    print("No matching files found.")


# %%
#input processo
processo = int(input("Número de processo: "))
#get row
try:
    row_index = df.index[df['Processo'] == processo].tolist()
    print(row_index)
    row = df.iloc[row_index[0]]
except:
    print('Processo não encontrado.')
    input('Enter para sair')

# %%
doc = Document("template.docx")

# %%
def identificacao(doc, row):
    for paragraph in doc.paragraphs:
        # Replace the placeholder for 'Raça'
        if '<<raca>>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<<raca>>', str(row['Raça']).lower())
        # Replace the placeholder for 'Sexo'
        if '<<sexo>>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<<sexo>>', 'sexo ' + str(row['Sexo']).lower())
        # Replace the placeholder for 'Profissão'
        if '<<profissao>>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<<profissao>>', str(row['Profissão']).lower())
        if '<<idade>>' in paragraph.text:
            #Obter idade
            nascimento = row['Data de Nascimento']
            hoje = row['Hora de início']
            idade = hoje.year - nascimento.year
            if (hoje.month, hoje.day) < (nascimento.month, nascimento.day):
                idade -= 1
            paragraph.text = paragraph.text.replace('<<idade>>', str(idade)+' anos')
    print('Identificacao')
    return doc

# %%
def antecedentes(doc,row):
    doencas = ['Patologia Ocular', 'Patologia Neurodegenerativa', 'Patologia Cerebelosa', 'Patologia desmielinizante', 'Patologia vascular']
    ap = []
    for doenca in doencas:
        if pd.isna(row[doenca]):
            row[doenca] = 'não;'
        if row[doenca] != 'não;':
            ap1 = row[doenca].split(';')
            #Remove empty strings or 'nao' string, if the user chose any option but also nao by mistake
            cleaned_ap1 = [s.strip() for s in ap1 if s.strip() and s.strip() != "não"]
            #Capitalize first letter
            cleaned_ap1 = [s[0].upper() + s[1:] if s else '' for s in cleaned_ap1]  
            ap = ap + cleaned_ap1
    doencas_bin = ['Polineuropatia', 'Mielopatia', 'Hipotensão ortostática','Exposição a fármacos ototóxicos']
    for doenca in doencas_bin:
        if row[doenca] == 'sim':
            ap.append(doenca)

    doenca = 'Outros antecedentes'
    cleaned_out = []
    if pd.notna(row[doenca]):
        cleaned_out = row[doenca].split(';')
        cleaned_out = [s.strip() if s else '' for s in cleaned_out]
        cleaned_out = [s[0].upper() + s[1:] if s else '' for s in cleaned_out]  
    ap = ap + cleaned_out  
    ap = '\n'.join(ap)         
    if not ap:
        ap = 'Nega'
    for paragraph in doc.paragraphs:
        if '<<antecedentes>>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<<antecedentes>>', ap)
    print('Antecedentes')
    return doc
    

# %%
def medicacao(doc,row):
    med = row['Medicação habitual2']
    if pd.notna(med):
        med = med.split(';')
        med = [s.strip() if s else '' for s in med]
        med = [s[0].upper() + s[1:] if s else '' for s in med]
        med = '\n'.join(med) 
    else:
        med = 'Nega'
    for paragraph in doc.paragraphs:
        if '<<medicacao>>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<<medicacao>>', med)
    print('Medicacao')
    return doc

# %%
def hda(doc,row):
    # Definir sintomas
    sintomas = row['Sintomas']
    if pd.notna(sintomas):
        sintomas = sintomas.split(';')
        sintomas = [s.strip() if s else '' for s in sintomas]
        sintomas = [s for s in sintomas if s.strip()]
        sintomas = [s.lower() if s else '' for s in sintomas]
        sintomas[0] = sintomas[0].capitalize()
        if len(sintomas)>1: 
            sintomas[-2] = sintomas[-2]+' e '
            sintomas [-2] = sintomas[-2]+sintomas[-1]
            sintomas = sintomas[:-1]
            sintomas = ', '.join(sintomas) 
        else:
            sintomas = sintomas[0]
    else:
        sintomas = 'Nega sintomas'
    
    # Definir duracao
    if pd.notna(row['Desde quando?']):
        unidades = row['Desde quando?']
        if unidades == 'dias':
            tempo = str(int(row['Quantos dias?']))+ ' '
        if unidades == 'semanas':
            tempo = str(int(row['Quantas semanas?']))+ ' '
        if unidades == 'meses':
            tempo = str(int(row['Quantos meses?']))+ ' '
        if unidades == 'anos':
            tempo = str(int(row['Quantos anos?']))+ ' '
        duracao = ' há ' + tempo + unidades
    else:
        duracao = ''    

    hda = sintomas + duracao + '.'
    for paragraph in doc.paragraphs:
        if '<<hda>>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<<hda>>', hda)
    print('Sintomas')
    return doc

# %%
def hda2(doc,row):
    carater = row['Caráter']
    if pd.isna(carater):
        carater = ''
    if carater == 'episódico':
        if pd.notna(row['Se episódico, duração']):
            unidades = row['Se episódico, duração']
            if unidades == 'segundos':
                tempo = str(int(row['Se episódico, quantos segundos']))+ ' '
            if unidades == 'minutos':
                tempo = str(int(row['Se episódico, quantos minutos']))+ ' '
            if unidades == 'horas':
                tempo = str(int(row['Se episódico, quantas horas']))+ ' '
            if unidades == 'dias':
                tempo = str(int(row['Se episódico, quantos dias']))+ ' '
            duracao = ' com duração de ' + tempo + unidades
        else:
            duracao = ''
    else:
        duracao = ''
    carater = carater + duracao
    descricao = ' Os sintomas têm um caráter ' + carater + '.'
    
    for paragraph in doc.paragraphs:
        if '<<hda2>>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<<hda2>>', descricao)

    print('Carater')
    return doc

# %%
def hda3(doc,row):
    apresentacao = row['Apresentação']
    if apresentacao == 'provocada':
        trigger = row['Se provocada, desencadeante:']
        if pd.notna(trigger):
            trigger = trigger.split(';')
            trigger = [s.strip() if s else '' for s in trigger]
            trigger = [s for s in trigger if s.strip()]
            trigger = [s.lower() if s else '' for s in trigger]
            if len(trigger)>1: 
                trigger[-2] = trigger[-2]+' e '
                trigger [-2] = trigger[-2]+trigger[-1]
                trigger = trigger[:-1]
                trigger = ', '.join(trigger) 
                print(trigger)
            else:
                trigger = trigger[0]
            descricao = ' Com apresentação provocada e apresentando como fator desencadeante: ' + trigger + '.'     
        else:
            descricao = ' Não consegue determinar fatores desencadeantes.' 
    if apresentacao == 'espontânea':
        descricao = ' Com apresentação espontânea.'
    else:
        descricao = ' Não se estabelecendo se apresentação espontânea ou provocada.'

    for paragraph in doc.paragraphs:
        if '<<hda3>>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<<hda3>>', descricao)

    print('Apresentacao')
    return doc

# %%
def hda4(doc,row):
    #Frase para sintomas neurologicos acompanhantes
    neurologicos = row['Sintomas acompanhantes neurológicos']
    if neurologicos != 'não;' and pd.notna(neurologicos):
        neurologicos = neurologicos.split(';')
        neurologicos = [s.strip() for s in neurologicos if s.strip() and s.strip() != "não"]
        neurologicos = [s for s in neurologicos if s.strip()]
        neurologicos = [s.lower() if s else '' for s in neurologicos]
        if len(neurologicos)>1: 
            neurologicos[-2] = neurologicos[-2]+' e '
            neurologicos [-2] = neurologicos[-2]+neurologicos[-1]
            neurologicos = neurologicos[:-1]
            neurologicos = ', '.join(neurologicos) 
        else:
            neurologicos = neurologicos[0]
        neurologicos = '\nComo sintomas neurológicos acompanhantes apresenta ' + neurologicos +'.'

    else:
        neurologicos = '\nSem sintomas neurológicos acompanhantes.'
    
    #Frase para sintomas otologicos acompanhantes
    otologicos = row['Sintomas acompanhantes otológicos']
    if otologicos != 'não;'and pd.notna(otologicos):
        otologicos = otologicos.split(';')
        otologicos = [s.strip() for s in otologicos if s.strip() and s.strip() != "não"]
        otologicos = [s for s in otologicos if s.strip()]
        otologicos = [s.lower() if s else '' for s in otologicos]
        if len(otologicos)>1: 
            otologicos[-2] = otologicos[-2]+' e '
            otologicos [-2] = otologicos[-2]+otologicos[-1]
            otologicos = otologicos[:-1]
            otologicos = ', '.join(otologicos) 
        else:
            otologicos = otologicos[0]
        otologicos = ' Como sintomas otológicos acompanhantes apresenta ' + otologicos +'.'

    else:
        otologicos = ' Sem sintomas otológicos acompanhantes.'
    
    #Frase para sintomas acompanhantes (outros)
    outros = row['Outros sintomas acompanhantes']
    if outros != 'não;'and pd.notna(outros):
        outros = outros.split(';')
        outros = [s.strip() for s in outros if s.strip() and s.strip() != "não"]
        outros = [s for s in outros if s.strip()]
        outros = [s.lower() if s else '' for s in outros]
        if len(outros)>1: 
            outros[-2] = outros[-2]+' e '
            outros [-2] = outros[-2]+outros[-1]
            outros = outros[:-1]
            outros = ', '.join(outros) 
        else:
            outros = outros[0]
        outros = ' Associadamente, apresenta ainda ' + outros +'.'

    else:
        outros = ''

    acompanhantes = neurologicos + otologicos + outros
    for paragraph in doc.paragraphs:
        if '<<hda4>>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<<hda4>>', acompanhantes)

    print('Acompanhantes')
    return doc

# %%
def en1(doc,row):
    #Campimetria
    campimetria = row['Campimetria']
    if campimetria == 'alterações mono-oculares':
        campimetria = 'Na avaliação campimétrica com alteraçõoes mono-oculares do olho ' + row['Se alterações monoculares, que olho?'] + '.'
    
    else:
        if pd.isna(campimetria):
            campimetria = "sem défices por confrontação"
        campimetria = 'Na avaliação campimétrica, ' + str(campimetria) + '.'
    #Pupilas
    pupilas = row['Anisocoria?'] 
    if pupilas == 'não' or pd.isna(pupilas):
        pupilas = ' Sem anisocoria.'
    else:
        pupilas = ' Anisocoria com ' + str(pupilas) + '.'
    
    #DPAR
    dpar = row['DPAR?'] 
    if dpar == 'não' or pd.isna(dpar):
        dpar = ' Sem DPAR.'
    else:
        dpar = ' DPAR no ' + str(dpar) + '.'

    # Movimentos oculares
    mov = row['Movimentos oculares']
    if pd.isna(mov):
        mov = 'preservados'
    mov = ' Movimentos oculares ' + str(mov) + '.'
    mov2 = row['Se movimentos oculares alterados, especificar parésia']
    if pd.isna(mov2):
        mov2 = ''
    else:
        mov2 = ' ' + str(mov2).capitalize() + '.'
    
    mov = str(mov) + str(mov2)
    
    # Sensibilidade na face
    sensibilidade = row['Sensibilidade álgica na face']
    if sensibilidade == 'alterada':
        sensibilidade = ' Sensibilidade da face alterada: ' + str(row['Se sensibilidade álgica na face alterada, explicitar território:']) + '.'
    else:
        sensibilidade = ' Sensibilidade da face sem alterações.'
    
    # Mimica
    mimica = row['Mímica facial']
    if mimica == 'alterada':
        mimica = mimica + str(row['Lateralidade da parésia facial']) + '.'
    else:
        mimica = ' Mímica facial sem alterações.'

    # Uvula
    uvula = row['Úvula']
    if pd.isna(uvula):
        uvula = 'centrada'
    uvula = ' Úvula ' + uvula + '.'

    # Lingua
    lingua = row['Protrusão da língua']
    if pd.isna(lingua):
        lingua = 'na linha média'
    lingua = ' Língua ' + lingua + '.'

    pares = campimetria + pupilas + dpar + mov + sensibilidade +  mimica + uvula + lingua
    for paragraph in doc.paragraphs:
        if '<<en1>>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<<en1>>', pares)


    print('Exame Neurologico: pares cranianos')
    return doc

# %%
def en2(doc,row):
    #Força global
    
    # Força global pbe
    pbe = row['Força global - prova de braços estendidos']
    if pd.isna(pbe):
        pbe = 'sem quedas ou desvios'
    if pbe != 'sem quedas ou desvios':
        pbe = pbe + ', ' + str(row['Prova de braços estendidos alterada (especificar):']).lower() + '.'
        pbe = 'Prova de braços estendidos alterada: ' + pbe +'.'
    else:
        pbe = 'PBE sem queda ou desvios.'
    pbe

    # Força global minga
    mingazzin = row['Força global - Mingazzini:']
    if pd.isna(mingazzin):
        mingazzin = 'sem quedas ou desvios'
    if mingazzin != 'sem quedas ou desvios':
        mingazzin = mingazzin + ', ' + str(row['Mingazzini alterado (especificar):']).lower() + '.'
        mingazzin = ' Prova de Mingazzini alterada: ' + mingazzin +'.'
    else:
        mingazzin = ' Mingazzini sem queda ou desvios. '
    mingazzin

    # Força segmentar
    seg = row['Força segmentar:']
    if pd.isna(seg):
        seg = 'normal nos 4 membros'
    if seg == 'normal nos 4 membros':
        seg = 'Força segmentar ' + seg +'.'
    else:
        if pd.notna(row['Força segmentar alterada (especifcar):']):
            seg = 'Força segmentar alterada: ' + str(row['Força segmentar alterada (especifcar):']) + '.'
        else:
            seg = 'Força segmentar alterada.'

    # ROT
    rot = row['ROTs']
    if rot == 'normais':
        rot = 'ROTs normais.'
    else:
        if pd.notna(row['ROTs alterados (especificar):']):
            rot = 'ROTs alterados: ' + str(row['ROTs alterados (especificar):']) + '.'
        else:
            rot = 'ROTs não avaliados.'

    # rcp
    rcp = row['Reflexos cutâneo plantares']
    if pd.isna(rcp):
        rcp = 'RCP não avaliado.'
    else:
        rcp = 'RCP com ' + rcp + '.'

    #Dismetria
    dismetria = row['Dismetria']
    if pd.isna(dismetria):
        dismetria = 'Dismetria ausente.'
    else:
        dismetria = 'Dismetria ' + dismetria + '.'
    
    forca = str(pbe) + str(mingazzin) + '\n' + str(seg) + '\n' + str(rot) + '\n' + str(rcp) + '\n' + str(dismetria)

    for paragraph in doc.paragraphs:
        if '<<en2>>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<<en2>>', forca)

    print('Exame Neurológico: Motricidade')
    return doc

# %%
def en3(doc,row):
    # Proprioceptiva

    prop = row['Sensibilidade proprioceptiva']
    if prop == 'alterada':
        prop = 'Sensibilidade proprioceptiva alterada: ' + str(row['Propiopceção alterada (onde?):']) + '.'
    else:
        prop = 'Sensibilidade proprioceptiva sem alterações.'

    # Fukuda
    fuk = row['Teste de Fukuda-Unterberger']
    if pd.isna(fuk):
        fuk = 'Não avaliado'
    fuk = 'Teste de Fukuda-Unterberger: ' + str(fuk) + '.'
    fuk

    # Romberg
    rom = row['Teste de Romberg']
    if pd.isna(rom):
        rom = 'Não avaliado'
    rom = 'Teste de Romberg: ' + str(rom) + '.'
    rom
    
    #Marcha
    marc = row['Marcha']
    if pd.notna(marc):
        marc = marc.split(';')
        marc = [s.strip() for s in marc if s.strip()]
        marc = [s for s in marc if s.strip()]
        marc = [s.lower() if s else '' for s in marc]
        if len(marc)>1: 
            marc[-2] = marc[-2]+' e '
            marc [-2] = marc[-2]+marc[-1]
            marc = marc[:-1]
            marc = ', '.join(marc) 
        else:
            marc = marc[0]
        marc = 'Marcha ' + marc +'.'

    else:
        marc = 'Marcha incaraterística.'

    sensibilidade = prop + '\n' + fuk + '\n' + rom + '\n' + marc

    for paragraph in doc.paragraphs:
        if '<<en3>>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<<en3>>', sensibilidade)

    print('Exame Neurológico: Sensibilidade')
    return doc

# %%
def nis(doc,row):
    
    nis_ger = row['Existe algum tipo de nistagmo (com fixação, sem fixação, evocado pelo olhar)?']
    if pd.isna(nis_ger):
        nis_ger =  'Não'
    if nis_ger == 'Não':
        nis_ger = 'Sem nistagmo com fixação, sem fixação ou evocado pelo olhar'
    if nis_ger == 'Sim':
        if row['Nistagmo espontâneo com fixação'] == 'sim':
            esquerda = str(row['Esquerda (velocidade média da fase lenta)'])
            direita = str(row['Direita (velocidade média da fase lenta)'])
            upbeat = str(row['Upbeat (velocidade média da fase lenta)'])
            downbeat = str(row['Downbeat (velocidade média da fase lenta)'])
            nis_espc = '\n'+'Nistagmo espontâneo COM fixação (velocidade média da fase lenta):'\
                +'\nEsquerda: ' + esquerda\
                +'\nDireita: ' + direita\
                +'\nUpbeat: ' + upbeat\
                +'\nDownbeat: ' + downbeat\
                +'\n'
        else:
            nis_espc = ''        

        if row['Nistagmo espontâneo sem fixação'] == 'sim':
            esquerda = str(row['Esquerda (velocidade média da fase lenta)2'])
            direita = str(row['Direita (velocidade média da fase lenta)2'])
            upbeat = str(row['Upbeat (velocidade média da fase lenta)2'])
            downbeat = str(row['Downbeat (velocidade média da fase lenta)2'])
            nis_esps = '\n'+'Nistagmo espontâneo SEM fixação (velocidade média da fase lenta):'\
                +'\nEsquerda: ' + esquerda\
                +'\nDireita: ' + direita\
                +'\nUpbeat: ' + upbeat\
                +'\nDownbeat: ' + downbeat\
                +'\n'

        else:
            nis_esps = ''    

        if row['Nistagmo evocado pelo olhar (posição primária)'] == 'sim':
            esquerda = str(row['Esquerda (velocidade média da fase lenta)3'])
            direita = str(row['Direita (velocidade média da fase lenta)3'])
            upbeat = str(row['Upbeat (velocidade média da fase lenta)3'])
            downbeat = str(row['Downbeat (velocidade média da fase lenta)3'])
            nis_evp = '\n'+'Nistagmo evocado pelo olhar (posição primária):'\
                +'\nEsquerda: ' + esquerda\
                +'\nDireita: ' + direita\
                +'\nUpbeat: ' + upbeat\
                +'\nDownbeat: ' + downbeat\
                +'\n'

        else:
            nis_evp = ''    

        if row['Nistagmo evocado pelo olhar (direita)'] == 'sim':
            esquerda = str(row['Esquerda (velocidade média da fase lenta)4'])
            direita = str(row['Direita (velocidade média da fase lenta)4'])
            upbeat = str(row['Upbeat (velocidade média da fase lenta)4'])
            downbeat = str(row['Downbeat (velocidade média da fase lenta)4'])
            nis_evd = '\n'+'Nistagmo evocado pelo olhar (direita):'\
                +'\nEsquerda: ' + esquerda\
                +'\nDireita: ' + direita\
                +'\nUpbeat: ' + upbeat\
                +'\nDownbeat: ' + downbeat\
                +'\n'

        else:
            nis_evd = ''   

        if row['Nistagmo evocado pelo olhar (esquerda)'] == 'sim':
            esquerda = str(row['Esquerda (velocidade média da fase lenta)5'])
            direita = str(row['Direita (velocidade média da fase lenta)5'])
            upbeat = str(row['Upbeat (velocidade média da fase lenta)5'])
            downbeat = str(row['Downbeat (velocidade média da fase lenta)5'])
            nis_eve = '\n'+'Nistagmo evocado pelo olhar (esquerda):'\
                +'\nEsquerda: ' + esquerda\
                +'\nDireita: ' + direita\
                +'\nUpbeat: ' + upbeat\
                +'\nDownbeat: ' + downbeat\
                +'\n'

        else:
            nis_eve = ''   

        if row['Nistagmo evocado pelo olhar (cima)'] == 'sim':
            esquerda = str(row['Esquerda (velocidade média da fase lenta)6'])
            direita = str(row['Direita (velocidade média da fase lenta)6'])
            upbeat = str(row['Upbeat (velocidade média da fase lenta)6'])
            downbeat = str(row['Downbeat (velocidade média da fase lenta)6'])
            nis_evc = '\n'+ 'Nistagmo evocado pelo olhar (cima):'\
                +'\nEsquerda: ' + esquerda\
                +'\nDireita: ' + direita\
                +'\nUpbeat: ' + upbeat\
                +'\nDownbeat: ' + downbeat\
                +'\n'

        else:
            nis_evc = ''   

        if row['Nistagmo evocado pelo olhar (baixo)'] == 'sim':
            esquerda = str(row['Esquerda (velocidade média da fase lenta)7'])
            direita = str(row['Direita (velocidade média da fase lenta)7'])
            upbeat = str(row['Upbeat (velocidade média da fase lenta)7'])
            downbeat = str(row['Downbeat (velocidade média da fase lenta)7'])
            nis_evb = '\n'+'Nistagmo evocado pelo olhar (baixo):'\
                +'\nEsquerda: ' + esquerda\
                +'\nDireita: ' + direita\
                +'\nUpbeat: ' + upbeat\
                +'\nDownbeat: ' + downbeat\
                +'\n'

        else:
            nis_evb = ''   
    
    if nis_ger == 'Sem nistagmo com fixação, sem fixação ou evocado pelo olhar':
        nis_ger = nis_ger
    else:
        nis_ger = 'Com nistagmo presente (com fixação, sem fixação ou evocado pelo olhar)' + nis_espc + nis_esps + nis_evp + nis_evd + nis_eve + nis_evc + nis_evb

    for paragraph in doc.paragraphs:
        if '<<nis1>>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<<nis1>>', nis_ger)

    print('Nistagmo')
    return doc

# %%
def per(doc,row):
    hg = row['Horizontal (Ganho)']
    hs = row['Horizontal (Simetria)']
    vg = row['Vertical (Ganho)']
    vs = row['Vertical (Simetria)']

    if pd.isna(hg) and pd.isna(hs) and pd.isna(vg) and pd.isna(vs):
        per = 'Perseguição lenta não avaliada. \n'
    else:
        per = 'Perseguição lenta:\n'
        if pd.isna(hg):
            hg = 'Horizontal (Ganho): não avaliado.\n'
        else:
            hg = 'Horizontal (Ganho): ' + str(hg) + '\n'

        if pd.isna(hs):
            hs = 'Horizontal (Simetria): não avaliado.\n'
        else:
            hs = 'Horizontal (Simetria): ' + str(hs) + '\n'

        if pd.isna(vg):
            vg = 'Vertical (Ganho): não avaliado.\n'
        else:
            vg = 'Vertical (Ganho): ' + str(vg) + '\n'

        if pd.isna(vs):
            vs = 'Vertical (Simetria): não avaliado.\n'
        else:
            vs = 'Vertical (Simetria): ' + str(vs) + '\n'

        per = per + hg + hs + vg + vs 

    for paragraph in doc.paragraphs:
        if '<<per>>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<<per>>', per)

    print('Perseguição lenta')
    return doc

# %%
def sac(doc,row):
    shv = row['Velocidade sacadas aleatórias horizontais ']
    shp = row['Precisão sacadas aleatórias horizontais ']
    shl = row['Latência sacadas aleatórias verticais']
    
    svv = row['Velocidade sacadas aleatórias verticais']
    svp = row['Precisão sacadas aleatórias verticais']
    svl = row['Latência sacadas aleatórias verticais']
    
    if pd.isna(shv) and pd.isna(shp) and pd.isna(shl) and pd.isna(svv) and pd.isna(svp) and pd.isna(svl):
        sac = 'Sacadas aleatórias não avaliadas. \n'
    else:
        sac = 'Sacadas aleatórias:\n'

        if pd.isna(shv):
            shv = 'Horizontal (Velocidade): não avaliado.\n'
        else:
            shv = 'Horizontal (Velocidade): ' + str(shv) + '\n'

        if pd.isna(shp):
            shp = 'Horizontal (Precisão: não avaliado.\n'
        else:
            shp = 'Horizontal (Precisão): ' + str(shp) + '\n'

        if pd.isna(shl):
            shl = 'Horizontal (Latência): não avaliado.\n'
        else:
            shl = 'Horizontal (Latência): ' + str(shl) + '\n\n'

        if pd.isna(svv):
            svv = 'Vertical (Velocidade): não avaliado.\n'
        else:
            svv = 'Vertical (Velocidade): ' + str(svv) + '\n'

        if pd.isna(svp):
            svp = 'Vertical (Precisão: não avaliado.\n'
        else:
            svp = 'Vertical (Precisão): ' + str(svp) + '\n'

        if pd.isna(svl):
            svl = 'Vertical (Latência): não avaliado.\n'
        else:
            svl = 'Vertical (Latência): ' + str(svl) + '\n'

        sac = sac + shv + shp + shl + svv + svp + svl

    for paragraph in doc.paragraphs:
        if '<<sac>>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<<sac>>', sac)

    print('Sacadas aleatórias')
    return doc

# %%
def opto(doc,row):
    o20d = row['Ganho 20º Direita']
    o40d = row['Ganho 40º Direita']
    o20e = row['Ganho 20º Esquerda']
    o40e = row['Ganho 40º Esquerda']
    
    if pd.isna(o20d) and pd.isna(o40d) and pd.isna(o20e) and pd.isna(o40e):
        opto = 'Nistagmo optocinético não avaliado. \n'
    else:
        opto = 'Nistagmo optocinético:\n'

        if pd.isna(o20d):
            o20d = 'Ganho 20º Direita: não avaliado.\n'
        else:
            o20d = 'Ganho 20º Direita: ' + str(o20d) + '\n'

        if pd.isna(o40d):
            o40d = 'Ganho 40º Direita: não avaliado.\n'
        else:
            o40d = 'Ganho 40º Direita: ' + str(o40d) + '\n'

        if pd.isna(o20e):
            o20e = 'Ganho 20º Esquerda: não avaliado.\n'
        else:
            o20e = 'Ganho 20º Esquerda: ' + str(o20e) + '\n'

        if pd.isna(o40e):
            o40e = 'Ganho 40º Esquerda: não avaliado.\n'
        else:
            o40e = 'Ganho 40º Esquerda: ' + str(o40e) + '\n'

        opto = opto + o20d + o40d + o20e + o40e

    for paragraph in doc.paragraphs:
        if '<<opto>>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<<opto>>', opto)

    print('Nistagmo Optocinético')
    return doc

# %%
def dix(doc,row):
    dix = row['Alguma alteração no Dix Hallpike?']
    if pd.isna(dix):
        dix = 'Dix-Hallpike: não realizado.\n'
        dixd = ''
        dixed = ''
    if dix == 'não':
        dix = 'Dix-Hallpike: sem alterações.\n'
        dixd = ''
        dixed = ''
    if dix == 'sim':
        dix = 'Dix-Hallpike: alterado.\n\n'
        dir = 'À direita:\n'
        esq = 'À esquerda:\n'
        #Dix-Hallpike à direita
        dixd = row['Nistagmo no Dix Hallpike para a direita em decúbito?']
        if pd.isna(dixd):
            dixd = dir + 'Não realizado em decúbito.\n'
        if dixd == 'não':
            dixd = dir + 'Sem alterações em decúbito.\n'
        if dixd == 'sim':
            dixd = dir + 'Alterado em decúbito.\n'

            dixde = row['Esquerda (velocidade média da fase lenta)8']
            dixdd = row['Direita (velocidade média da fase lenta)8']
            dixdc = row['Upbeat (velocidade média da fase lenta)8']
            dixdb = row['Downbeat (velocidade média da fase lenta)8']
            if pd.isna(dixde) and pd.isna(dixdd) and pd.isna(dixdc) and pd.isna(dixdb):
                dixd = dixd
            else:
                if pd.isna(dixde):
                    dixde = 'Esquerda (velocidade média da fase lenta): não avaliada\n'
                else:
                    dixde = 'Esquerda (velocidade média da fase lenta):' + str(dixde) + '\n'

                if pd.isna(dixdd):
                    dixdd = 'Direita (velocidade média da fase lenta): não avaliada\n'
                else:
                    dixdd = 'Direita (velocidade média da fase lenta):' + str(dixdd) + '\n'
                if pd.isna(dixdc):
                    dixdc = 'Upbeat (velocidade média da fase lenta): não avaliada\n'
                else:
                    dixdc = 'Upbeat (velocidade média da fase lenta):' + str(dixdc) + '\n'
                if pd.isna(dixdb):
                    dixdb = 'Downbeat (velocidade média da fase lenta): não avaliada\n\n'
                else:
                    dixdb = 'Downbeat (velocidade média da fase lenta):' + str(dixdb) + '\n\n'

                dixd = dixd + dixdd + dixde + dixdc + dixdb
        dixs = row['Nistagmo no Dix Hallpike para a direita sentado?']
        if pd.isna(dixs):
            dixs = 'Não realizado sentado.\n'
        if dixs == 'não':
            dixs = 'Sem alterações sentado.\n'
        if dixs == 'sim':
            dixs = 'Alterado sentado.\n'

            dixse = row['Esquerda (velocidade média da fase lenta)9']
            dixsd = row['Direita (velocidade média da fase lenta)9']
            dixsc = row['Upbeat (velocidade média da fase lenta)9']
            dixsb = row['Downbeat (velocidade média da fase lenta)9']
            if pd.isna(dixse) and pd.isna(dixsd) and pd.isna(dixsc) and pd.isna(dixsb):
                dixs = dixs
            else:
                if pd.isna(dixse):
                    dixse = 'Esquerda (velocidade média da fase lenta): não avaliada\n'
                else:
                    dixse = 'Esquerda (velocidade média da fase lenta):' + str(dixse) + '\n'

                if pd.isna(dixsd):
                    dixsd = 'Direita (velocidade média da fase lenta): não avaliada\n'
                else:
                    dixsd = 'Direita (velocidade média da fase lenta):' + str(dixsd) + '\n'
                if pd.isna(dixsc):
                    dixsc = 'Upbeat (velocidade média da fase lenta): não avaliada\n'
                else:
                    dixsc = 'Upbeat (velocidade média da fase lenta):' + str(dixsc) + '\n'
                if pd.isna(dixsb):
                    dixsb = 'Downbeat (velocidade média da fase lenta): não avaliada\n'
                else:
                    dixsb = 'Downbeat (velocidade média da fase lenta):' + str(dixsb) + '\n'

                dixs = dixs + dixsd + dixse + dixsc + dixsb
        dixd = dixd + dixs + '\n'
        #Dix Hallpike à esquerda
        dixed = row['Nistagmo no Dix Hallpike para a esquerda decúbito?']
        if pd.isna(dixed):
            dixed = esq + 'Não realizado em decúbito.\n'
        if dixed == 'não':
            dixed = esq + 'Sem alterações em decúbito.\n'
        if dixed == 'sim':
            dixed = esq + 'Alterado em decúbito.\n'

            dixede = row['Esquerda (velocidade média da fase lenta)10']
            dixedd = row['Direita (velocidade média da fase lenta)10']
            dixedc = row['Upbeat (velocidade média da fase lenta)10']
            dixedb = row['Downbeat (velocidade média da fase lenta)10']
            if pd.isna(dixede) and pd.isna(dixedd) and pd.isna(dixedc) and pd.isna(dixedb):
                dixed = dixed
            else:
                if pd.isna(dixede):
                    dixede = 'Esquerda (velocidade média da fase lenta): não avaliada\n'
                else:
                    dixede = 'Esquerda (velocidade média da fase lenta):' + str(dixede) + '\n'

                if pd.isna(dixedd):
                    dixedd = 'Direita (velocidade média da fase lenta): não avaliada\n'
                else:
                    dixedd = 'Direita (velocidade média da fase lenta):' + str(dixedd) + '\n'
                if pd.isna(dixedc):
                    dixedc = 'Upbeat (velocidade média da fase lenta): não avaliada\n'
                else:
                    dixedc = 'Upbeat (velocidade média da fase lenta):' + str(dixedc) + '\n'
                if pd.isna(dixedb):
                    dixedb = 'Downbeat (velocidade média da fase lenta): não avaliada\n\n'
                else:
                    dixedb = 'Downbeat (velocidade média da fase lenta):' + str(dixedb) + '\n\n'

                dixed = dixed + dixedd + dixede + dixedc + dixedb
        dixes = row['Nistagmo no Dix Hallpike para a esquerda sentado?']
        if pd.isna(dixes):
            dixes = 'Não realizado sentado.\n'
        if dixes == 'não':
            dixes = 'Sem alterações sentado.\n'
        if dixes == 'sim':
            dixes = 'Alterado sentado.\n'

            dixese = row['Esquerda (velocidade média da fase lenta)11']
            dixesd = row['Direita (velocidade média da fase lenta)11']
            dixesc = row['Upbeat (velocidade média da fase lenta)11']
            dixesb = row['Downbeat (velocidade média da fase lenta)11']
            if pd.isna(dixese) and pd.isna(dixesd) and pd.isna(dixesc) and pd.isna(dixesb):
                dixes = dixes
            else:
                if pd.isna(dixese):
                    dixese = 'Esquerda (velocidade média da fase lenta): não avaliada\n'
                else:
                    dixese = 'Esquerda (velocidade média da fase lenta):' + str(dixese) + '\n'

                if pd.isna(dixesd):
                    dixesd = 'Direita (velocidade média da fase lenta): não avaliada\n'
                else:
                    dixesd = 'Direita (velocidade média da fase lenta):' + str(dixesd) + '\n'
                if pd.isna(dixesc):
                    dixesc = 'Upbeat (velocidade média da fase lenta): não avaliada\n'
                else:
                    dixesc = 'Upbeat (velocidade média da fase lenta):' + str(dixesc) + '\n'
                if pd.isna(dixesb):
                    dixesb = 'Downbeat (velocidade média da fase lenta): não avaliada\n'
                else:
                    dixesb = 'Downbeat (velocidade média da fase lenta):' + str(dixesb) + '\n'

                dixes = dixes + dixesd + dixese + dixesc + dixesb
        dixed = dixed + dixes + '\n'
    dix = dix + dixd + dixed     

    for paragraph in doc.paragraphs:
        if '<<dix>>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<<dix>>', dix)

    print('Dix-Hallpike')
    return doc

# %%
def pag(doc,row):
    pag = row['Alguma alteração no Pagnini Mc-Clure']
    if pd.isna(pag):
        pag = 'Pagnini Mc-Clure: não realizado.\n'
        pagd = ''
        paged = ''
    if pag == 'não':
        pag = 'Pagnini Mc-Clure: sem alterações.\n'
        pagd = ''
        paged = ''
    if pag == 'sim':
        pag = 'Pagnini Mc-Clure: alterado.\n\n'
        dir = 'À direita:\n'
        esq = 'À esquerda:\n'
        #Pagnini à direita
        pagd = row['Nistagmo no Pagnini Mc-Clure para a direita em decúbito?']
        if pd.isna(pagd):
            pagd = dir + 'Não realizado em decúbito.\n'
        if pagd == 'não':
            pagd = dir + 'Sem alterações em decúbito.\n'
        if pagd == 'sim':
            pagd = dir + 'Alterado em decúbito.\n'

            pagde = row['Esquerda (velocidade média da fase lenta)12']
            pagdd = row['Direita (velocidade média da fase lenta)12']
            pagdc = row['Upbeat (velocidade média da fase lenta)12']
            pagdb = row['Downbeat (velocidade média da fase lenta)12']
            if pd.isna(pagde) and pd.isna(pagdd) and pd.isna(pagdc) and pd.isna(pagdb):
                pagd = pagd
            else:
                if pd.isna(pagde):
                    pagde = 'Esquerda (velocidade média da fase lenta): não avaliada\n'
                else:
                    pagde = 'Esquerda (velocidade média da fase lenta):' + str(pagde) + '\n'

                if pd.isna(pagdd):
                    pagdd = 'Direita (velocidade média da fase lenta): não avaliada\n'
                else:
                    pagdd = 'Direita (velocidade média da fase lenta):' + str(pagdd) + '\n'
                if pd.isna(pagdc):
                    pagdc = 'Upbeat (velocidade média da fase lenta): não avaliada\n'
                else:
                    pagdc = 'Upbeat (velocidade média da fase lenta):' + str(pagdc) + '\n'
                if pd.isna(pagdb):
                    pagdb = 'Downbeat (velocidade média da fase lenta): não avaliada\n\n'
                else:
                    pagdb = 'Downbeat (velocidade média da fase lenta):' + str(pagdb) + '\n\n'

                pagd = pagd + pagdd + pagde + pagdc + pagdb
        
        pags = row['Nistagmo no Pagnini Mc-Clure para a direita sentado?']
        if pd.isna(pags):
            pags = 'Não realizado sentado.\n'
        if pags == 'não':
            pags = 'Sem alterações sentado.\n'
        if pags == 'sim':
            pags = 'Alterado sentado.\n'

            pagse = row['Esquerda (velocidade média da fase lenta)13']
            pagsd = row['Direita (velocidade média da fase lenta)13']
            pagsc = row['Upbeat (velocidade média da fase lenta)13']
            pagsb = row['Downbeat (velocidade média da fase lenta)13']
            if pd.isna(pagse) and pd.isna(pagsd) and pd.isna(pagsc) and pd.isna(pagsb):
                pags = pags
            else:
                if pd.isna(pagse):
                    pagse = 'Esquerda (velocidade média da fase lenta): não avaliada\n'
                else:
                    pagse = 'Esquerda (velocidade média da fase lenta):' + str(pagse) + '\n'

                if pd.isna(pagsd):
                    pagsd = 'Direita (velocidade média da fase lenta): não avaliada\n'
                else:
                    pagsd = 'Direita (velocidade média da fase lenta):' + str(pagsd) + '\n'
                if pd.isna(pagsc):
                    pagsc = 'Upbeat (velocidade média da fase lenta): não avaliada\n'
                else:
                    pagsc = 'Upbeat (velocidade média da fase lenta):' + str(pagsc) + '\n'
                if pd.isna(pagsb):
                    pagsb = 'Downbeat (velocidade média da fase lenta): não avaliada\n'
                else:
                    pagsb = 'Downbeat (velocidade média da fase lenta):' + str(pagsb) + '\n'

                pags = pags + pagsd + pagse + pagsc + pagsb
        pagd = pagd + pags + '\n'
        #Pagnini à esquerda
        paged = row['Nistagmo no Pagnini Mc-Clure para a esquerda em decúbito?']
        if pd.isna(paged):
            paged = esq + 'Não realizado em decúbito.\n'
        if paged == 'não':
            paged = esq + 'Sem alterações em decúbito.\n'
        if paged == 'sim':
            paged = esq + 'Alterado em decúbito.\n'

            pagede = row['Esquerda (velocidade média da fase lenta)14']
            pagedd = row['Direita (velocidade média da fase lenta)14']
            pagedc = row['Upbeat (velocidade média da fase lenta)14']
            pagedb = row['Downbeat (velocidade média da fase lenta)14']
            if pd.isna(pagede) and pd.isna(pagedd) and pd.isna(pagedc) and pd.isna(pagedb):
                paged = paged
            else:
                if pd.isna(pagede):
                    pagede = 'Esquerda (velocidade média da fase lenta): não avaliada\n'
                else:
                    pagede = 'Esquerda (velocidade média da fase lenta):' + str(pagede) + '\n'

                if pd.isna(pagedd):
                    pagedd = 'Direita (velocidade média da fase lenta): não avaliada\n'
                else:
                    pagedd = 'Direita (velocidade média da fase lenta):' + str(pagedd) + '\n'
                if pd.isna(pagedc):
                    pagedc = 'Upbeat (velocidade média da fase lenta): não avaliada\n'
                else:
                    pagedc = 'Upbeat (velocidade média da fase lenta):' + str(pagedc) + '\n'
                if pd.isna(pagedb):
                    pagedb = 'Downbeat (velocidade média da fase lenta): não avaliada\n\n'
                else:
                    pagedb = 'Downbeat (velocidade média da fase lenta):' + str(pagedb) + '\n\n'

                paged = paged + pagedd + pagede + pagedc + pagedb
        pages = row['Nistagmo no Pagnini Mc-Clure para a esquerda sentado?']
        if pd.isna(pages):
            pages = 'Não realizado sentado.\n'
        if pages == 'não':
            pages = 'Sem alterações sentado.\n'
        if pages == 'sim':
            pages = 'Alterado sentado.\n'

            pagese = row['Esquerda (velocidade média da fase lenta)15']
            pagesd = row['Direita (velocidade média da fase lenta)15']
            pagesc = row['Upbeat (velocidade média da fase lenta)15']
            pagesb = row['Downbeat (velocidade média da fase lenta)15']
            if pd.isna(pagese) and pd.isna(pagesd) and pd.isna(pagesc) and pd.isna(pagesb):
                pages = pages
            else:
                if pd.isna(pagese):
                    pagese = 'Esquerda (velocidade média da fase lenta): não avaliada\n'
                else:
                    pagese = 'Esquerda (velocidade média da fase lenta):' + str(pagese) + '\n'

                if pd.isna(pagesd):
                    pagesd = 'Direita (velocidade média da fase lenta): não avaliada\n'
                else:
                    pagesd = 'Direita (velocidade média da fase lenta):' + str(pagesd) + '\n'
                if pd.isna(pagesc):
                    pagesc = 'Upbeat (velocidade média da fase lenta): não avaliada\n'
                else:
                    pagesc = 'Upbeat (velocidade média da fase lenta):' + str(pagesc) + '\n'
                if pd.isna(pagesb):
                    pagesb = 'Downbeat (velocidade média da fase lenta): não avaliada\n'
                else:
                    pagesb = 'Downbeat (velocidade média da fase lenta):' + str(pagesb) + '\n'

                pages = pages + pagesd + pagese + pagesc + pagesb
        paged = paged + pages + '\n'
    pag = pag + pagd + paged     

    for paragraph in doc.paragraphs:
        if '<<pag>>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<<pag>>', pag)

    print('Pagnini Mc-Clure')
    return doc

# %%
def posic(doc,row):
    posic = row['Há nistagmo posicional?']
    if pd.isna(posic):
        posic = '\nNistagmo posicional não avaliado.'
    if posic == 'sim':
        # Nistagmo posicional em supuino
        poss = row['Nistagmo posicional em supino?']
        if pd.isna(poss):
            poss = '\nNistagmo posicional em supino não avaliado.'
        if poss == 'sim':
            esquerda = str(row['Esquerda (velocidade média da fase lenta)16'])
            direita = str(row['Direita (velocidade média da fase lenta)16'])
            upbeat = str(row['Upbeat (velocidade média da fase lenta)16'])
            downbeat = str(row['Downbeat (velocidade média da fase lenta)16'])
            poss = '\n'+'Com nistagmo posicional em supino:'\
            +'\nEsquerda: ' + esquerda\
            +'\nDireita: ' + direita\
            +'\nUpbeat: ' + upbeat\
            +'\nDownbeat: ' + downbeat\
            +'\n'
        if poss == 'não':
            poss = '\nSem nistagmo posicional em supino.'

        # Nistagmo posiciional à direita
        posd = row['Nistagmo posicional à direita2']
        if pd.isna(posd):
            posd = '\nNistagmo posicional à direita não avaliado.'
        if posd == 'sim':
            esquerda = str(row['Esquerda (velocidade média da fase lenta)17'])
            direita = str(row['Direita (velocidade média da fase lenta)17'])
            upbeat = str(row['Upbeat (velocidade média da fase lenta)17'])
            downbeat = str(row['Downbeat (velocidade média da fase lenta)17'])
            posd = '\n'+'Com nistagmo posicional à direita:'\
            +'\nEsquerda: ' + esquerda\
            +'\nDireita: ' + direita\
            +'\nUpbeat: ' + upbeat\
            +'\nDownbeat: ' + downbeat\
            +'\n'
        if posd == 'não':
            posd = '\nSem nistagmo posicional à direita.'
    
        # Nistagmo posiciional à esquerda
        pose = row['Nistagmo posicional à esquerda2']
        if pd.isna(pose):
            pose = '\nNistagmo posicional à esquerda não avaliado.'
        if pose == 'sim':
            esquerda = str(row['Esquerda (velocidade média da fase lenta)18'])
            direita = str(row['Direita (velocidade média da fase lenta)18'])
            upbeat = str(row['Upbeat (velocidade média da fase lenta)18'])
            downbeat = str(row['Downbeat (velocidade média da fase lenta)18'])
            pose = '\n'+'Com nistagmo posicional à esquerda:'\
            +'\nEsquerda: ' + esquerda\
            +'\nDireita: ' + direita\
            +'\nUpbeat: ' + upbeat\
            +'\nDownbeat: ' + downbeat\
            +'\n'
        if pose == 'não':
            pose = '\nSem nistagmo posicional à esquerda.'

        posic = poss + posd + pose


    if posic == 'não':
        posic = 'Sem nistagmo posicional.\n'


    for paragraph in doc.paragraphs:
        if '<<posic>>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<<posic>>', posic)

    print('Nistagmo posicional')
    return doc

# %%
def posic(doc,row):
    posic = row['Há nistagmo posicional?']
    if pd.isna(posic):
        posic = '\nNistagmo posicional não avaliado.'
    if posic == 'sim':
        # Nistagmo posicional em supuino
        poss = row['Nistagmo posicional em supino?']
        if pd.isna(poss):
            poss = '\nNistagmo posicional em supino não avaliado.'
        if poss == 'sim':
            esquerda = str(row['Esquerda (velocidade média da fase lenta)16'])
            direita = str(row['Direita (velocidade média da fase lenta)16'])
            upbeat = str(row['Upbeat (velocidade média da fase lenta)16'])
            downbeat = str(row['Downbeat (velocidade média da fase lenta)16'])
            poss = '\n'+'Com nistagmo posicional em supino:'\
            +'\nEsquerda: ' + esquerda\
            +'\nDireita: ' + direita\
            +'\nUpbeat: ' + upbeat\
            +'\nDownbeat: ' + downbeat\
            +'\n'
        if poss == 'não':
            poss = '\nSem nistagmo posicional em supino.'

        # Nistagmo posiciional à direita
        posd = row['Nistagmo posicional à direita2']
        if pd.isna(posd):
            posd = '\nNistagmo posicional à direita não avaliado.'
        if posd == 'sim':
            esquerda = str(row['Esquerda (velocidade média da fase lenta)17'])
            direita = str(row['Direita (velocidade média da fase lenta)17'])
            upbeat = str(row['Upbeat (velocidade média da fase lenta)17'])
            downbeat = str(row['Downbeat (velocidade média da fase lenta)17'])
            posd = '\n'+'Com nistagmo posicional à direita:'\
            +'\nEsquerda: ' + esquerda\
            +'\nDireita: ' + direita\
            +'\nUpbeat: ' + upbeat\
            +'\nDownbeat: ' + downbeat\
            +'\n'
        if posd == 'não':
            posd = '\nSem nistagmo posicional à direita.'
    
        # Nistagmo posiciional à esquerda
        pose = row['Nistagmo posicional à esquerda2']
        if pd.isna(pose):
            pose = '\nNistagmo posicional à esquerda não avaliado.'
        if pose == 'sim':
            esquerda = str(row['Esquerda (velocidade média da fase lenta)18'])
            direita = str(row['Direita (velocidade média da fase lenta)18'])
            upbeat = str(row['Upbeat (velocidade média da fase lenta)18'])
            downbeat = str(row['Downbeat (velocidade média da fase lenta)18'])
            pose = '\n'+'Com nistagmo posicional à esquerda:'\
            +'\nEsquerda: ' + esquerda\
            +'\nDireita: ' + direita\
            +'\nUpbeat: ' + upbeat\
            +'\nDownbeat: ' + downbeat\
            +'\n'
        if pose == 'não':
            pose = '\nSem nistagmo posicional à esquerda.'

        posic = poss + posd + pose


    if posic == 'não':
        posic = 'Sem nistagmo posicional.\n'


    for paragraph in doc.paragraphs:
        if '<<posic>>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<<posic>>', posic)

    print('Nistagmo posicional')
    return doc

# %%
def vhit(doc,row):
    vhid = {}
    vhid['adg'] = row['CSC AD - Ganho']
    vhid['adc'] = row['CSC AD']
    vhid['ado'] = row['CSC AD2'] 
    vhid['aeg'] = row['CSC AE - Ganho']
    vhid['aec'] = row['CSC AE']
    vhid['aeo'] = row['CSC AE2']
    vhid['hdg'] = row['CSC HD - Ganho']
    vhid['hdc'] = row['CSC HD']
    vhid['hdo'] = row['CSC HD2']
    vhid['heg'] = row['CSC HE - Ganho']
    vhid['hec'] = row['CSC HE']
    vhid['heo'] = row['CSC HE2']
    vhid['pdg'] = row['CSC PD - Ganho']
    vhid['pdc'] = row['CSC PD']
    vhid['pdo'] = row['CSC PD2']
    vhid['peg'] = row['CSC PE - Ganho']
    vhid['pec'] = row['CSC PE']
    vhid['peo'] = row['CSC PE2']

    for i in vhid.keys():
        vhid[i] = str(vhid[i])

    valid = 0
    for i in vhid.values():
        if pd.notna(i):
            valid = valid + 1
    if valid == 0:
        text = 'vHIT não realizado.'

    else:
        titulos = ['CSC Anterior Direito:',\
                   '\nCSC Anterior Esquerdo:',\
                   '\nCSC Horizontal Direito:',\
                   '\nCSC Horizontal Esquerdo:',\
                   '\nCSC Posterior Direito:',\
                   '\nCSC Posterior Esquerdo:']

        opcoes = ['\nGanho: ','\nSacadas covert: ','\nSacadas overt: ']

        text = ''
        i = 0
        for titulo in titulos:
            text = text + titulo
            for opcao in opcoes:
                text = text + opcao + list(vhid.values())[i]
                i = i+1
            text = text + '\n'
    for paragraph in doc.paragraphs:
        if '<<vhit>>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<<vhit>>', text)

    print('vHIT')
    return doc

# %%
def mcdt(doc,row):
    mcd2 = ''
    mcd3 = ''

    mcd1 = row['Exames realizados']
    if pd.isna(mcd1):
        mcd1 = 'Não'
    if mcd1 == 'Não':
        mcd1 = 'Sem exames realizados.'
    if mcd1 == 'Ambas':
        mcd1 = 'TC-CE e RM-CE já realizados.'
    if mcd1 == 'RM-CE':
        mcd1 = 'Realizou RM-CE.'
    if mcd1 == 'TC-CE':
        mcd1 = 'Realizou TC-CE.'

    if mcd1 != 'Sem exames realizados':
        mcd2 = row['Lesões']
        mcd3 = row['Localização']
        if pd.isna(mcd2) and pd.isna(mcd3):
            mcd2 = ' Sem lesões detetadas.'
        else:
            mcd2 = str(mcd2).lower()
            if pd.notna(mcd3):
                mcd3 = mcd3.split(';')
                mcd3 = [s.strip() if s else '' for s in mcd3]
                mcd3 = [s for s in mcd3 if s.strip()]
                mcd3 = [s.lower() if s else '' for s in mcd3]
                if len(mcd3)>1: 
                    mcd3[-2] = mcd3[-2]+' e '
                    mcd3[-2] = mcd3[-2]+mcd3[-1]
                    mcd3 = mcd3[:-1]
                    mcd3 = ', '.join(mcd3) 
                else:
                    mcd3 = mcd3[0]
            else:
                mcd3 = 'sem localização especificada.'

    text = mcd1 + ' Com lesão/ões ' + mcd2 + ' do ' + mcd3 +'.'
    for paragraph in doc.paragraphs:
        if '<<mcdt>>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<<mcdt>>', text)

    print('MCDTs')
    return doc

# %%
doc = identificacao(doc,row)
doc = antecedentes(doc,row)
doc = medicacao(doc,row)
doc = hda(doc,row)
doc = hda2(doc,row)
doc = hda3(doc,row)
doc = hda4(doc,row)
doc = en1(doc,row)
doc = en2(doc,row)
doc = en3(doc,row)
doc = nis(doc,row)
doc = per(doc,row)
doc = sac(doc, row)
doc = opto(doc,row)
doc = dix(doc,row)
doc = pag(doc,row)
doc = posic(doc,row)
doc = vhit(doc,row)
doc = mcdt(doc,row)

# %%
try:
    output_filename = f"output_{processo}.docx"
    doc.save(output_filename)
    print(f"Saved: {output_filename}")
except:
    print('Não foi possível gerar o ficheiro. Vê se não está já aberto')
    input('Enter para continuar')

# %%
# Open word file

# Create an instance of the Word application
word = win32.Dispatch("Word.Application")

# Make Word visible to the user
word.Visible = True

# Optionally, open a specific document (replace with your file path)
all = current_directory + '\\' + output_filename
doc = word.Documents.Open(all)
word.Selection.WholeStory()
word.Selection.Copy()


# %%
print(' ')
print(' ')

print('É SÓ COLAR NO SCLINICO :D')
print('Se não abriu um word ou deu algum erro, experimenta umas das opções:')
print(' - Garante que não havia nenhum word gerado por este programa já aberto antes')
print(' - Garante que só há um documento de excel com as respostas do forms na pasta')
print('Enter para sair')
input()


